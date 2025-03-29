import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
import sys

def create_conference_styles(document):
    # Helper function to get or create a style
    def get_or_create_style(style_name, style_type):
        if style_name in document.styles:
            return document.styles[style_name]
        return document.styles.add_style(style_name, style_type)
    
    # Title Style
    title_style = get_or_create_style('Conference Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(8)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(0)
    
    # Page Number Style
    page_num_style = get_or_create_style('Page Number', WD_STYLE_TYPE.PARAGRAPH)
    page_num_style.font.name = 'Times New Roman'
    page_num_style.font.size = Pt(8)
    page_num_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Session Title Style
    session_title_style = get_or_create_style('Session Title', WD_STYLE_TYPE.PARAGRAPH)
    session_title_style.font.name = 'Times New Roman'
    session_title_style.font.size = Pt(8)
    session_title_style.font.bold = True
    session_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author Style
    author_style = get_or_create_style('Author', WD_STYLE_TYPE.PARAGRAPH)
    author_style.font.name = 'Times New Roman'
    author_style.font.size = Pt(8)
    author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_style.paragraph_format.space_after = Pt(0)
    
    # Abstract Style
    abstract_style = get_or_create_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
    abstract_style.font.name = 'Times New Roman'
    abstract_style.font.size = Pt(8)
    abstract_style.paragraph_format.space_before = Pt(8)
    abstract_style.paragraph_format.left_indent = Inches(0.5)
    abstract_style.paragraph_format.right_indent = Inches(0.5)
    abstract_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Body Text Style
    body_text_style = get_or_create_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    body_text_style.font.name = 'Times New Roman'
    body_text_style.font.size = Pt(8)
    body_text_style.paragraph_format.space_after = Pt(8)
    body_text_style.paragraph_format.space_before = Pt(0)
    
    # References Label Style
    references_label_style = get_or_create_style('References Label', WD_STYLE_TYPE.PARAGRAPH)
    references_label_style.font.name = 'Times New Roman'
    references_label_style.font.size = Pt(8)
    references_label_style.font.bold = True
    references_label_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # references_label_style.paragraph_format.space_before = Pt(8)
    
    # Reference Entry Style
    reference_entry_style = get_or_create_style('Reference Entry', WD_STYLE_TYPE.PARAGRAPH)
    reference_entry_style.font.name = 'Times New Roman'
    reference_entry_style.font.size = Pt(8)
    # reference_entry_style.paragraph_format.first_line_indent = Inches(0)
    # reference_entry_style.paragraph_format.left_indent = Inches(0.5)
    # reference_entry_style.paragraph_format.hanging_indent = Inches(0.5)
    # reference_entry_style.paragraph_format.space_after = Pt(6)
    
    return document

def create_conference_session(document, title, author_str, abstract, proposal, references):
    # Add conference title
    # p = document.add_paragraph('Conference on Higher Education Pedagogy', 'Conference Title')
    
    # Add page number (you might want to handle this differently, e.g., in headers/footers)
    # p = document.add_paragraph('3', 'Page Number')
    
    # Add session title
    p = document.add_heading(title, level=1)
    # p = document.add_paragraph(title, 'Session Title')
    
    # Add authors with italicized affiliations
    for line in author_str.split("\n"):
        if "*" in line:
            parts = line.split(", *")
            authors = parts[0]
            institution = parts[1].strip("*")
            p = document.add_paragraph(authors, 'Author')
            p.add_run(f", {institution}").italic = True
        else:
            document.add_paragraph(line, 'Author')
    
    # Add abstract
    p = document.add_paragraph('', 'Abstract')
    p.add_run('Abstract: ').bold = True
    p.add_run(abstract).bold = False
    p.style = 'Abstract'

    # Add the proposal
    cleaned_proposal = [para.strip() for para in proposal if para.strip()]
    for para in cleaned_proposal:
        p = document.add_paragraph(para, 'Body Text')
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
    
    # # Add references
    if references:
        p = document.add_paragraph('References', 'References Label')
        for ref in references:
            p = document.add_paragraph(ref, 'Reference Entry')
    
    return document

def create_authors(authors, affiliations):
    # Split authors and affiliations into lists
    author_list = authors.split(", ")
    affiliation_list = affiliations.split(", ")

    # Create a dictionary to map superscripts to affiliations
    affiliation_dict = {}
    for aff in affiliation_list:
        if aff[0].isdigit():  # Check if the first character is a superscript
            key = aff[0]
            value = aff[1:].strip()
            affiliation_dict[key] = value

    # Group authors by their corresponding affiliations
    institution_authors = {}
    for author in author_list:
        if any(char.isdigit() for char in author):  # Check for superscripts
            superscript = ''.join(filter(str.isdigit, author))
            author_name = ''.join(filter(lambda x: not x.isdigit(), author)).strip()
            institution = affiliation_dict.get(superscript, "")
            if institution not in institution_authors:
                institution_authors[institution] = []
            institution_authors[institution].append(author_name)
        else:
            if "Unknown" not in institution_authors:
                institution_authors["Unknown"] = []
            institution_authors["Unknown"].append(author.strip())

    # Format the output with authors grouped by institution
    formatted_output = []
    for institution, authors in institution_authors.items():
        authors_str = ", ".join(authors)
        formatted_output.append(f"{authors_str}, *{institution}*")

    return "\n".join(formatted_output)

if __name__ == '__main__':
    try:
        # Check if the CSV file exists
        csv_file = 'chep_data.csv'
        if not os.path.exists(csv_file):
            raise FileNotFoundError(f"CSV file '{csv_file}' not found. Please ensure it is in the same directory as the script.")

        # Load the CSV file
        df = pd.read_csv(csv_file)

        # Check if required columns are present
        required_columns = ['Submission title', 'Submission authors', 'Affiliations', 'Abstract', 'Proposal', 'References']
        for column in required_columns:
            if column not in df.columns:
                raise ValueError(f"Missing required column '{column}' in the CSV file.")

        # Alphabetize by title
        df = df.sort_values(by='Submission title')

        # Create the document
        document = Document()

        # Create the conference styles
        document = create_conference_styles(document)

        # Iterate over each row in the CSV
        for index, row in df.iterrows():
            # Extract the relevant information from the row
            title = row['Submission title']
            authors = row['Submission authors']
            affiliations = row['Affiliations']
            abstract = row['Abstract']
            proposal = row['Proposal']
            references = row['References']

            # Ensure proposal is a string and handle missing values and newlines
            proposal = str(proposal) if not pd.isna(proposal) else ""
            # Split the proposal into paragraphs
            proposal = proposal.replace("\r", "").split("\n\n")

            # Ensure references is a string and handle missing values
            references = str(references) if not pd.isna(references) else ""
            # Split the references into a list
            references = references.split("\n")
            # For each reference, remove the leading number and period and any whitespace
            references = [
                ref.split(". ", 1)[1].strip() if ". " in ref else ref.strip()
                for ref in references
            ]
            # Remove any empty strings, trailing, or leading whitespace
            references = [ref.strip() for ref in references if ref]

            # Create the author string with formatted affiliations
            author_str = create_authors(authors, affiliations)

            # Create a new conference session in the document
            document = create_conference_session(document, title, author_str, abstract, proposal, references)

        # Save the document to a DOCX file
        document.save('booklet.docx')
        print("Booklet created successfully: 'booklet.docx'")

    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        sys.exit(1)
